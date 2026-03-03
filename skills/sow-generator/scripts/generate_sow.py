#!/usr/bin/env python3
"""SOW Generator — processes a .docx template with find-and-replace and conditional logic.

Usage:
    python generate_sow.py --template INPUT.docx --output OUTPUT.docx --config '{"key": "val"}'

The config JSON determines which replacements and conditional logic to apply.
See REFERENCE.md for placeholder and conditional logic details.
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph, handling text split across runs.

    Returns the number of replacements made (0 or 1).
    """
    if old_text not in paragraph.text:
        return 0

    # Try single-run replacement first (preserves formatting best)
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return 1

    # Placeholder spans multiple runs — join, replace, put in first run
    runs = paragraph.runs
    if runs:
        full_text = ''.join(r.text for r in runs)
        if old_text in full_text:
            runs[0].text = full_text.replace(old_text, new_text)
            for r in runs[1:]:
                r.text = ''
            return 1

    return 0


def replace_in_table(table, old_text, new_text):
    """Replace text in all cells of a table. Returns count of replacements."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                count += replace_in_paragraph(paragraph, old_text, new_text)
    return count


def delete_paragraph(paragraph):
    """Remove a paragraph from the document by deleting its XML element."""
    p = paragraph._element
    p.getparent().remove(p)


def delete_paragraphs_containing(doc, search_text):
    """Delete all paragraphs whose text contains search_text. Returns count deleted."""
    to_delete = [p for p in doc.paragraphs if search_text in p.text]
    for p in to_delete:
        delete_paragraph(p)
    return len(to_delete)


def delete_section(doc, start_text, end_text):
    """Delete all paragraphs from one containing start_text through one containing end_text.

    Inclusive of both boundary paragraphs.
    """
    deleting = False
    to_delete = []
    for para in doc.paragraphs:
        if not deleting and start_text in para.text:
            deleting = True
        if deleting:
            to_delete.append(para)
        if deleting and end_text in para.text:
            break
    for para in to_delete:
        delete_paragraph(para)
    return len(to_delete)


def apply_conditional_logic(doc, config):
    """Apply template-specific conditional removals.

    Args:
        doc: The Document object (modified in place).
        config: The config dict with template type and field values.
    """
    template_type = config.get('template', 'migration')

    if template_type == 'migration':
        if not config.get('discovery', False):
            # After replacements, #platform is already replaced with the actual
            # platform name, so use the config value for the search string.
            platform = config.get('platform', '#platform')
            if platform.lower() in ("don't know", "unknown", "idk", "?", "tbd"):
                platform = '#platform'
            # Delete the deliverables reference line FIRST — it also contains
            # "Exhibit 2: Discovery & Extraction Services" text, which would
            # cause delete_section to start too early and wipe out Exhibit 1.
            delete_paragraphs_containing(
                doc, f'If the source {platform} system requires database discovery'
            )
            delete_section(doc, 'Exhibit 2: Discovery & Extraction Services', '[END]')

    elif template_type == 'avid':
        # Remove server rental note if client-provided hypervisor
        hypervisor = config.get('hypervisor', '').lower()
        if hypervisor in ('client vmware', 'client proxmox'):
            delete_paragraphs_containing(doc, 'Server rental fees')

        # Remove network caveats if 10Gb
        network = config.get('network', '').lower()
        if network == '10gb':
            caveat_fragments = [
                '1Gb networking is supported; however',
                'NSA recommends bonded 1Gb connections',
                'All NSA-provided systems are equipped with multiple',
            ]
            for fragment in caveat_fragments:
                delete_paragraphs_containing(doc, fragment)


def process_document(template_path, output_path, replacements):
    """Open a .docx template, apply replacements, save to output.

    Args:
        template_path: Path to the source .docx file.
        output_path: Path to save the processed .docx file.
        replacements: Dict of {old_text: new_text} to apply.

    Returns:
        Dict with 'replacements_made' (int) and 'not_found' (list of keys with 0 matches).
    """
    doc = Document(template_path)
    total = 0
    not_found = []

    for old_text, new_text in replacements.items():
        count = 0

        # Replace in body paragraphs
        for paragraph in doc.paragraphs:
            count += replace_in_paragraph(paragraph, old_text, new_text)

        # Replace in tables
        for table in doc.tables:
            count += replace_in_table(table, old_text, new_text)

        # Replace in headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and header.is_linked_to_previous is False:
                    for paragraph in header.paragraphs:
                        count += replace_in_paragraph(paragraph, old_text, new_text)
                    for table in header.tables:
                        count += replace_in_table(table, old_text, new_text)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer.is_linked_to_previous is False:
                    for paragraph in footer.paragraphs:
                        count += replace_in_paragraph(paragraph, old_text, new_text)
                    for table in footer.tables:
                        count += replace_in_table(table, old_text, new_text)

        if count == 0:
            not_found.append(old_text)
        total += count

    doc.save(output_path)
    return {'replacements_made': total, 'not_found': not_found}


def main():
    parser = argparse.ArgumentParser(description='Process SOW template')
    parser.add_argument('--template', required=True, help='Path to .docx template')
    parser.add_argument('--output', required=True, help='Output .docx path')
    parser.add_argument('--config', required=True, help='JSON config string')
    args = parser.parse_args()

    config = json.loads(args.config)
    template_type = config.get('template', 'migration')

    # Build replacement map based on template type
    replacements = {}
    placeholder_map = {
        'client': '#client',
        'reseller': '#reseller',
        'date': '#date',
    }

    if template_type == 'migration':
        placeholder_map['platform'] = '#platform'
    elif template_type == 'avid':
        placeholder_map.update({
            'destination': '#destination',
            'total_data_size': '#total_data_size',
            'disk_data': '#disk_data',
            'clip_count': '#clip_count',
            'lto_data': '#lto_data',
            'archive_system': '#archive_system',
            'free_storage': '#free_storage',
        })

    for key, placeholder in placeholder_map.items():
        value = config.get(key)
        if value and value.lower() not in ("don't know", "unknown", "idk", "?", "tbd"):
            replacements[placeholder] = value

    # Handle logo — always replace with client name as placeholder text
    if 'client' in config:
        replacements['#logo'] = config['client']

    # Handle Avid-specific computed replacements
    if template_type == 'avid':
        hypervisor = config.get('hypervisor', '')
        hypervisor_map = {
            'client vmware': 'Client-provided VMWare',
            'client proxmox': 'Client-provided Proxmox',
            'nsa proxmox': 'NSA-provided Proxmox hardware',
        }
        hyp_replacement = hypervisor_map.get(hypervisor.lower())
        if hyp_replacement:
            replacements['#hypervisor'] = hyp_replacement

        network = config.get('network', '')
        network_map = {
            '10gb': '10Gb host bus adapter (HBA)',
            '1gb': '1Gb networking',
            'bonded 1gb': 'Bonded 1Gb connections',
        }
        net_replacement = network_map.get(network.lower())
        if net_replacement:
            replacements['#network'] = net_replacement

    result = process_document(args.template, args.output, replacements)

    # Apply conditional logic (section/line removal)
    doc = Document(args.output)
    apply_conditional_logic(doc, config)
    doc.save(args.output)

    print(json.dumps(result, indent=2))


if __name__ == '__main__':
    main()
